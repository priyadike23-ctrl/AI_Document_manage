"""
OnlyTech Portal - Backend API
Handles: Auth (signup/login), PlantMind AI (Groq), Documents (with real file storage
in Supabase Storage + real PDF/DOCX text extraction), AI-assisted before/after
document updates, Document Approval workflow WITH AUTOMATIC TITLE+CONTENT MATCH
ROUTING, Maintenance tickets, Activity log, and dashboard stats.

APPROVAL ROUTING (fully automatic, per the Decision Engine)
------------------------------------------------------------------
Every employee document upload (propose_upload) and every proposed document
replacement (propose_document_update) is:

  1. Text-extracted (PDF / DOCX / plain text)
  2. Matched against the existing Knowledge Base by TITLE similarity
  3. COMPARED against the matched document's actual content (content similarity)
  4. Combined into a single 0-100 match score
  5. Routed by the Decision Engine:

        score >= AI_AUTO_APPROVE_MIN_SCORE (default 90)  -> AUTO-PUBLISHED
        AI_AUTO_REJECT_MAX_SCORE <= score < APPROVE (70-89) -> ADMIN REVIEW QUEUE
        score <  AI_AUTO_REJECT_MAX_SCORE (default 70)    -> AUTO-REJECTED

  6. The outcome (Approved / Pending / Rejected) is always written to
     document_approvals for a full audit trail, and to activity_log so it
     shows up instantly in the admin Overview -> Recent Activity feed and
     the Approvals tab (only Pending rows are ever listed there).

Only the middle band ever requires a human click in
/api/approvals/<id>/decide, OR gets picked up by the batch "Automate" button
in /api/approvals/automate-batch, which re-runs the SAME Groq-powered
comparison (KB document vs proposed file) on every currently-Pending row.

GROQ-POWERED COMPARISON (used by both the single-upload flow AND the
"Automate" batch button)
------------------------------------------------------------------
Groq does not host an embeddings endpoint, so instead of a local
sentence-transformers model we ask the Groq-hosted LLM directly: "here is
the CURRENT published document, here is the PROPOSED replacement/upload -
how consistent/same-topic are they, on a scale of 0-100?" That score is
combined with a cheap, dependency-free title-match score (difflib) to
produce the final routing score. This means the automation genuinely uses
your GROQ_API_KEY end-to-end and needs no extra ML dependencies.

PLANTMIND AI KNOWLEDGE SCOPE (updated)
------------------------------------------------------------------
PlantMind AI (/api/ai/ask) is grounded on REAL data pulled live from
Supabase at question time - not just published Knowledge Base documents.
On every question it now searches:

  - documents            -> Active AND Draft/Updated documents (full text)
  - document_history     -> full version-by-version audit trail (who changed
                             what, when, and what the content looked like at
                             that version)
  - document_approvals   -> the Approvals queue itself: Pending / Approved /
                             Rejected requests, who submitted them, the AI
                             match score + reasoning, who decided them
  - maintenance_tickets  -> all incidents (open + resolved), resolution notes
  - activity_log         -> the full audit trail of who did what and when
                             (uploads, approvals, rejections, resolutions,
                             deletions, batch automation runs)
  - users                -> employee directory (name/department/role/email)

Every context block handed to Groq is built directly from a live Supabase
SELECT, so PlantMind AI can never invent a document, employee, or incident
that isn't actually in the database - it can only summarize/explain what's
really stored. Each answer is returned together with a `sources` array so
the admin UI can show exactly which records were used.

Run:
    pip install -r requirements.txt
    cp .env.example .env      # fill in your keys
    python app.py

requirements.txt:
    flask
    flask-cors
    python-dotenv
    werkzeug
    supabase
    groq
    pypdf
    python-docx
"""

import io
import os
import re
import json
import uuid
import difflib
import collections
import datetime
from functools import wraps

from flask import Flask, request, jsonify, Response
from flask_cors import CORS
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename

# ---- PDF / DOCX extraction libraries ---------------------------------------
PYPDF_AVAILABLE = False
DOCX_AVAILABLE = False
try:
    from pypdf import PdfReader
    PYPDF_AVAILABLE = True
except ImportError:
    print("[WARN] pypdf not installed. Run: pip install pypdf")

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    print("[WARN] python-docx not installed. Run: pip install python-docx")

# --------------------------------------------------------------------------
# ENV / CONFIG
# --------------------------------------------------------------------------
load_dotenv()

GROQ_API_KEY        = os.getenv("GROQ_API_KEY")
SUPABASE_URL         = os.getenv("SUPABASE_URL")
SUPABASE_KEY         = os.getenv("SUPABASE_KEY")   # service_role key for server-side writes
GROQ_MODEL_NAME      = os.getenv("GROQ_MODEL", "llama-3.3-70b-versatile")
STORAGE_BUCKET       = os.getenv("SUPABASE_STORAGE_BUCKET", "documents")

# Max characters of extracted document text we'll keep / send to Groq per document.
MAX_EXTRACTED_CHARS = 15000
# Max characters of each side (current vs proposed) sent to Groq for the
# comparison call specifically - kept smaller so the batch automation call
# stays fast and cheap even on large documents.
MAX_COMPARE_CHARS = 4000

# PlantMind AI / Decision Engine thresholds (0-100). These DRIVE the actual
# auto-publish / admin-review / auto-reject routing decision below.
AI_AUTO_APPROVE_MIN_SCORE = int(os.getenv("AI_AUTO_APPROVE_MIN_SCORE", "90"))
AI_AUTO_REJECT_MAX_SCORE  = int(os.getenv("AI_AUTO_REJECT_MAX_SCORE", "70"))

AUTO_APPROVER_NAME = "PlantMind AI (Auto-Approved)"
AUTO_REJECTOR_NAME = "PlantMind AI (Auto-Rejected)"

app = Flask(__name__)
CORS(app)

# ---- Supabase client -------------------------------------------------------
supabase = None
try:
    from supabase import create_client, Client
    if SUPABASE_URL and SUPABASE_KEY:
        supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)
    else:
        print("[WARN] SUPABASE_URL / SUPABASE_KEY not set yet.")
except ImportError:
    print("[WARN] supabase package not installed. Run: pip install supabase")

# ---- Groq client -----------------------------------------------------------
groq_client = None
try:
    from groq import Groq
    if GROQ_API_KEY:
        groq_client = Groq(api_key=GROQ_API_KEY)
    else:
        print("[WARN] GROQ_API_KEY not set yet.")
except ImportError:
    print("[WARN] groq package not installed. Run: pip install groq")


def require_supabase(fn):
    @wraps(fn)
    def wrapper(*args, **kwargs):
        if supabase is None:
            return jsonify({"error": "Supabase is not configured. Add SUPABASE_URL and SUPABASE_KEY to .env"}), 500
        return fn(*args, **kwargs)
    return wrapper


def now_iso():
    return datetime.datetime.utcnow().isoformat()


def pick_icon(filename_or_name):
    ext = filename_or_name.rsplit(".", 1)[-1].lower() if "." in filename_or_name else ""
    return {
        "pdf": "📕", "doc": "📝", "docx": "📝", "txt": "📝", "md": "📝",
        "xls": "📊", "xlsx": "📊", "csv": "📊",
        "png": "🖼️", "jpg": "🖼️", "jpeg": "🖼️", "gif": "🖼️", "webp": "🖼️",
        "zip": "🗜️", "json": "🧾", "log": "🧾",
    }.get(ext, "📄")


def log_activity(actor, action, target=None, details=None):
    """Best-effort write to activity_log; never blocks the main request on failure."""
    if supabase is None:
        return
    try:
        supabase.table("activity_log").insert({
            "id": str(uuid.uuid4()),
            "actor": actor,
            "action": action,
            "target": target,
            "details": details,
            "created_at": now_iso(),
        }).execute()
    except Exception as e:
        print("[WARN] activity log failed:", e)


def insert_approval_audit_row(record, ai_result):
    """
    Best-effort write of a Pending approval row into document_approvals,
    including ai_score / ai_reasoning columns if the table has them. Falls
    back to inserting without those two columns if the schema doesn't have
    them yet, so this never breaks the upload flow on an older database.
    """
    if supabase is None:
        return
    enriched = dict(record)
    enriched["ai_score"] = ai_result.get("score")
    enriched["ai_reasoning"] = ai_result.get("reasoning")
    try:
        supabase.table("document_approvals").insert(enriched).execute()
    except Exception:
        try:
            supabase.table("document_approvals").insert(record).execute()
        except Exception as e:
            print("[WARN] Could not write approval audit row:", e)


# --------------------------------------------------------------------------
# DECISION ENGINE HELPERS
# --------------------------------------------------------------------------
def decide_from_score(score):
    """Pure threshold routing used by the automatic decision engine."""
    if score >= AI_AUTO_APPROVE_MIN_SCORE:
        return "approve"
    if score < AI_AUTO_REJECT_MAX_SCORE:
        return "reject"
    return "review"


def publish_pending_record(record, decided_by):
    """
    Immediately publish a pending upload/update record to the Knowledge Base
    (used by the AUTO-APPROVE path, both the live upload flow and the batch
    "Automate" button). Writes document_history, an Approved audit row in
    document_approvals, and an activity_log entry - exactly like a manual
    admin approval would, just without the click. Returns the document_id.
    """
    if record.get("document_id"):
        doc_id = record["document_id"]
        updates = {
            "version": record["proposed_version"],
            "status": "Active",
            "content": record["new_content"],
            "updated_at": now_iso(),
        }
        if record.get("file_path"):
            updates.update({
                "file_path": record["file_path"],
                "file_name": record["file_name"],
                "file_type": record["file_type"],
                "file_size": record["file_size"],
            })
        supabase.table("documents").update(updates).eq("id", doc_id).execute()
    else:
        doc_id = str(uuid.uuid4())
        supabase.table("documents").insert({
            "id": doc_id,
            "name": record["doc_name"],
            "icon": pick_icon(record.get("file_name") or record["doc_name"]),
            "version": record["proposed_version"],
            "status": "Active",
            "summary": "",
            "content": record["new_content"],
            "file_path": record.get("file_path"),
            "file_name": record.get("file_name"),
            "file_type": record.get("file_type"),
            "file_size": record.get("file_size"),
            "uploaded_by": record["submitted_by"],
            "created_at": now_iso(),
            "updated_at": now_iso(),
        }).execute()

    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()), "document_id": doc_id, "version": record["proposed_version"],
        "updated_by": record["submitted_by"],
        "changes": f"Automatically published by {decided_by}.",
        "content_snapshot": record["new_content"], "created_at": now_iso(),
    }).execute()

    # If this record came from an existing document_approvals row (batch
    # automation path), update that row in place instead of inserting a
    # duplicate audit entry.
    if record.get("id"):
        try:
            res = supabase.table("document_approvals").update({
                "status": "Approved", "decided_at": now_iso(), "decided_by": decided_by,
            }).eq("id", record["id"]).execute()
            if not res.data:
                raise Exception("no matching row")
        except Exception:
            audit = dict(record)
            audit["status"] = "Approved"
            audit["decided_at"] = now_iso()
            audit["decided_by"] = decided_by
            try:
                supabase.table("document_approvals").insert(audit).execute()
            except Exception as e:
                print("[WARN] Could not write auto-approval audit row:", e)

    log_activity(decided_by, "approved_edit", record["doc_name"], f"Auto-published {record['proposed_version']}")
    return doc_id


def discard_pending_record(record, decided_by, reason=""):
    """
    Immediately discard an upload/update record (used by the AUTO-REJECT
    path, both the live upload flow and the batch "Automate" button).
    Deletes the uploaded file from storage and writes a Rejected audit row +
    activity_log entry so the decision is fully auditable even though
    nothing was ever published.
    """
    if record.get("file_path"):
        delete_from_storage(record["file_path"])

    if record.get("id"):
        try:
            res = supabase.table("document_approvals").update({
                "status": "Rejected", "decided_at": now_iso(), "decided_by": decided_by,
                "file_path": None,
            }).eq("id", record["id"]).execute()
            if not res.data:
                raise Exception("no matching row")
        except Exception:
            audit = dict(record)
            audit["status"] = "Rejected"
            audit["decided_at"] = now_iso()
            audit["decided_by"] = decided_by
            audit["file_path"] = None
            try:
                supabase.table("document_approvals").insert(audit).execute()
            except Exception as e:
                print("[WARN] Could not write auto-rejection audit row:", e)

    log_activity(decided_by, "rejected_edit", record["doc_name"], reason or "Automatically rejected - low match score")


# --------------------------------------------------------------------------
# TEXT EXTRACTION  ->  PDF / DOCX / plain text
# --------------------------------------------------------------------------
def extract_pdf_text(file_bytes):
    if not PYPDF_AVAILABLE:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for page in reader.pages:
            try:
                parts.append(page.extract_text() or "")
            except Exception:
                continue
        return "\n".join(parts).strip()
    except Exception as e:
        print("[WARN] PDF extraction failed:", e)
        return ""


def extract_docx_text(file_bytes):
    if not DOCX_AVAILABLE:
        return ""
    try:
        doc = DocxDocument(io.BytesIO(file_bytes))
        parts = [p.text for p in doc.paragraphs if p.text]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        parts.append(cell.text)
        return "\n".join(parts).strip()
    except Exception as e:
        print("[WARN] DOCX extraction failed:", e)
        return ""


def extract_text(file_bytes, mimetype, filename):
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mimetype = mimetype or ""

    try:
        if mimetype.startswith("text/") or ext in ("txt", "md", "csv", "json", "log"):
            text = file_bytes.decode("utf-8", errors="ignore")
        elif ext == "pdf" or mimetype == "application/pdf":
            text = extract_pdf_text(file_bytes)
        elif ext == "docx" or mimetype == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            text = extract_docx_text(file_bytes)
        else:
            text = ""
    except Exception as e:
        print("[WARN] extract_text failed:", e)
        text = ""

    if text and len(text) > MAX_EXTRACTED_CHARS:
        text = text[:MAX_EXTRACTED_CHARS]

    return text


def ensure_doc_content(doc):
    """Returns the document's text content, extracting + caching it from
    storage on first use if the `content` column is empty."""
    existing = (doc.get("content") or "").strip()
    if existing:
        return existing

    file_path = doc.get("file_path")
    if not file_path or supabase is None:
        return ""

    try:
        file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(file_path)
    except Exception as e:
        print(f"[WARN] Could not download {file_path} for extraction:", e)
        return ""

    text = extract_text(
        file_bytes,
        doc.get("file_type"),
        doc.get("file_name") or doc.get("name") or file_path,
    )

    if text:
        try:
            supabase.table("documents").update({"content": text}).eq("id", doc["id"]).execute()
        except Exception as e:
            print("[WARN] Could not cache extracted content:", e)

    return text


def ensure_approval_new_content(approval):
    """
    Returns the proposed/new text content for an approval row, extracting it
    from the attached file in Storage (and caching it back onto the row) if
    the `new_content` column is empty. Mirrors ensure_doc_content() but for
    document_approvals rows.
    """
    existing = (approval.get("new_content") or "").strip()
    if existing:
        return existing

    file_path = approval.get("file_path")
    if not file_path or supabase is None:
        return ""

    try:
        file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(file_path)
    except Exception as e:
        print(f"[WARN] Could not download {file_path} for extraction:", e)
        return ""

    text = extract_text(file_bytes, approval.get("file_type"), approval.get("file_name") or file_path)

    if text and approval.get("id"):
        try:
            supabase.table("document_approvals").update({"new_content": text}).eq("id", approval["id"]).execute()
        except Exception as e:
            print("[WARN] Could not cache extracted approval content:", e)

    return text


def upload_bytes_to_storage(file_bytes, mimetype, original_filename, folder="docs"):
    unique_path = f"{folder}/{uuid.uuid4().hex}_{secure_filename(original_filename)}"
    if supabase:
        supabase.storage.from_(STORAGE_BUCKET).upload(
            unique_path, file_bytes, {"content-type": mimetype or "application/octet-stream"}
        )
    return unique_path


def delete_from_storage(file_path):
    if not file_path or supabase is None:
        return
    try:
        supabase.storage.from_(STORAGE_BUCKET).remove([file_path])
    except Exception as e:
        print(f"[WARN] Could not delete {file_path} from storage:", e)


# --------------------------------------------------------------------------
# TITLE MATCHING  ->  cheap, dependency-free first pass used to find which
# Knowledge Base document (if any) a filename/doc_name is likely about.
# --------------------------------------------------------------------------
def normalize_title(name):
    if not name:
        return ""
    name = name.rsplit(".", 1)[0] if "." in name else name
    name = name.lower()
    name = re.sub(r"[^a-z0-9]+", " ", name).strip()
    return name


def find_best_title_match(candidate_name, existing_docs):
    norm_candidate = normalize_title(candidate_name)
    if not norm_candidate:
        return 0, None

    best_score = 0
    best_doc = None
    for d in existing_docs:
        norm_existing = normalize_title(d.get("name") or "")
        if not norm_existing:
            continue
        ratio = difflib.SequenceMatcher(None, norm_candidate, norm_existing).ratio() * 100
        if ratio > best_score:
            best_score = ratio
            best_doc = d
    return round(best_score), best_doc


def content_similarity(text_a, text_b, sample_chars=3000):
    """Cheap, dependency-free fallback similarity score (0-100) between two
    blocks of text. Only used when Groq is unavailable or errors out."""
    a = (text_a or "").strip().lower()[:sample_chars]
    b = (text_b or "").strip().lower()[:sample_chars]
    if not a or not b:
        return None
    return round(difflib.SequenceMatcher(None, a, b).ratio() * 100)


# --------------------------------------------------------------------------
# GROQ-POWERED CONTENT COMPARISON  ->  the actual "compare file from
# Knowledge Base against file from Approval" step, used by BOTH the
# single-upload Decision Engine and the batch "Automate" button.
# --------------------------------------------------------------------------
def groq_compare_documents(doc_name, current_text, proposed_text):
    """
    Asks Groq to score how consistent/same-topic two document excerpts are.
    Returns {"score": int 0-100, "reasoning": str}. Falls back to a
    dependency-free text-diff score if Groq is unavailable or errors.
    """
    current_text = (current_text or "").strip()
    proposed_text = (proposed_text or "").strip()

    if not current_text or not proposed_text:
        return {"score": 0, "reasoning": "One of the two documents had no extractable text to compare."}

    if groq_client is None:
        fallback = content_similarity(current_text, proposed_text) or 0
        return {"score": fallback, "reasoning": "Groq is not configured; used a plain text-similarity fallback."}

    system_prompt = """You are a document-comparison engine for OnlyTech's internal Knowledge Base.
You will be shown the CURRENT PUBLISHED VERSION of a document and a PROPOSED
REPLACEMENT / NEW UPLOAD. Decide how consistent they are in subject matter,
scope, and structure - i.e. is the proposed file clearly an updated version
of the same document (same procedures/policy/system, just revised), or is it
a substantially different, unrelated, low-quality, or placeholder file?

Score guide:
- 90-100: Same topic, clearly a legitimate updated version. Safe to auto-publish.
- 70-89: Related / plausible update but with notable gaps or differences. Needs a human look.
- 0-69: Different topic, inconsistent, low-quality, or junk content. Should be rejected.

Reply with ONLY a compact JSON object, nothing else - no markdown fences, no commentary:
{"score": <integer 0-100>, "reasoning": "<one or two plain sentences>"}"""

    user_content = (
        f"Document title: {doc_name}\n\n"
        f"--- CURRENT PUBLISHED VERSION (excerpt) ---\n{current_text[:MAX_COMPARE_CHARS]}\n\n"
        f"--- PROPOSED REPLACEMENT / UPLOAD (excerpt) ---\n{proposed_text[:MAX_COMPARE_CHARS]}"
    )

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_content},
            ],
            temperature=0.1,
            max_tokens=250,
        )
        raw = (completion.choices[0].message.content or "").strip()
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        parsed = json.loads(raw)
        score = int(parsed.get("score", 50))
        score = max(0, min(100, score))
        reasoning = str(parsed.get("reasoning", "")).strip()[:500] or "Groq evaluated this comparison automatically."
        return {"score": score, "reasoning": reasoning}
    except Exception as e:
        print("[WARN] groq_compare_documents failed:", e)
        fallback = content_similarity(current_text, proposed_text)
        if fallback is None:
            fallback = 0
        return {"score": fallback, "reasoning": f"Groq comparison failed ({e}); used a plain text-similarity fallback."}


def score_upload_by_title_match(filename, doc_name, content_text):
    """
    Runs the full "check Knowledge Base -> find title match -> Groq-compare
    content -> score" pipeline for a NEW employee upload and returns the
    routing decision for the Decision Engine:

        {"decision": "approve" | "review" | "reject", "score": int,
         "reasoning": str, "matched_document": dict | None}
    """
    existing_docs = []
    if supabase is not None:
        try:
            existing_docs = supabase.table("documents") \
                .select("id,name,version,content,status,file_path,file_type,file_name") \
                .eq("status", "Active").execute().data or []
        except Exception as e:
            print("[WARN] Could not fetch existing documents for title match:", e)

    title_score, matched_doc = find_best_title_match(doc_name or filename, existing_docs)

    content_score = None
    if matched_doc:
        existing_content = ensure_doc_content(matched_doc)
        groq_result = groq_compare_documents(doc_name or filename, existing_content, content_text)
        content_score = groq_result["score"]

    if content_score is not None:
        # Title tells us "is this probably the same document"; Groq's
        # content comparison confirms "does it actually say the same kind
        # of thing". Weight content higher since it's the more reliable
        # signal, but still require some title agreement.
        combined_score = round(title_score * 0.3 + content_score * 0.7)
    else:
        combined_score = title_score

    decision = decide_from_score(combined_score)
    # Never auto-approve if we don't actually have a document to publish
    # a new version against.
    if decision == "approve" and not matched_doc:
        decision = "review"

    if matched_doc:
        reasoning = (f'Title match {title_score}% against "{matched_doc["name"]}" '
                     f'({matched_doc.get("version", "v1.0")}); Groq content comparison '
                     f'{content_score if content_score is not None else "n/a"}%. Combined score: {combined_score}%.')
    else:
        reasoning = (f"No comparable document title was found in the Knowledge Base "
                     f"(best title match: {title_score}%).")

    return {
        "decision": decision,
        "score": combined_score,
        "reasoning": reasoning[:900],
        "matched_document": matched_doc,
    }


def score_approval_with_groq(approval):
    """
    Re-scores an EXISTING document_approvals row (used by the batch
    "Automate" button). This is the "compare file from Knowledge Base
    against file from Approval" step:

      1. Resolve which Knowledge Base document this approval is about -
         either the linked document_id (fast path, already resolved when
         the employee submitted it) or, if there is none (a brand-new
         upload with no target yet), fall back to a fresh title match.
      2. Pull the KB document's CURRENT text (ensure_doc_content) and the
         approval's PROPOSED text (ensure_approval_new_content), extracting
         from Supabase Storage on demand if either is missing.
      3. Ask Groq to score how consistent the two are (0-100).
      4. Combine with a title-match score exactly like the live upload path,
         so a batch re-score gives the same kind of number a fresh upload
         would have gotten.

    Returns: {"score": int, "reasoning": str, "matched_document": dict|None}
    """
    doc_name = approval.get("doc_name") or "Unknown Document"
    document_id = approval.get("document_id")
    proposed_text = ensure_approval_new_content(approval)

    matched_doc = None
    title_score = 0

    if document_id and supabase is not None:
        try:
            res = supabase.table("documents").select(
                "id,name,version,content,status,file_path,file_type,file_name"
            ).eq("id", document_id).execute()
            if res.data:
                matched_doc = res.data[0]
                title_score = 100  # explicitly linked at submission time
        except Exception as e:
            print("[WARN] Could not fetch linked document for approval:", e)

    if matched_doc is None:
        existing_docs = []
        if supabase is not None:
            try:
                existing_docs = supabase.table("documents").select(
                    "id,name,version,content,status,file_path,file_type,file_name"
                ).eq("status", "Active").execute().data or []
            except Exception as e:
                print("[WARN] Could not fetch documents for approval title match:", e)
        title_score, matched_doc = find_best_title_match(doc_name, existing_docs)

    if matched_doc is None:
        return {
            "score": title_score,
            "reasoning": f"No comparable document was found in the Knowledge Base (best title match: {title_score}%).",
            "matched_document": None,
        }

    current_text = ensure_doc_content(matched_doc) or (approval.get("old_content") or "")
    groq_result = groq_compare_documents(doc_name, current_text, proposed_text)
    content_score = groq_result["score"]

    combined_score = round(title_score * 0.3 + content_score * 0.7)
    reasoning = (f'Linked/matched to "{matched_doc["name"]}" ({matched_doc.get("version", "v1.0")}) '
                 f'- title match {title_score}%, Groq content comparison {content_score}% '
                 f'({groq_result["reasoning"]}). Combined score: {combined_score}%.')

    return {"score": combined_score, "reasoning": reasoning[:900], "matched_document": matched_doc}


# --------------------------------------------------------------------------
# AI TEXT ANALYSIS  ->  used by the "Update Docs" before/after dashboard
# --------------------------------------------------------------------------
def analyze_text_with_ai(doc_name, text):
    text = (text or "").strip()
    if not text:
        return {
            "summary": "No readable text could be extracted from this file "
                       "(likely an image, spreadsheet, or scanned document).",
            "key_points": [],
        }

    if groq_client is None:
        return {"summary": "PlantMind AI is not configured on the server.", "key_points": []}

    system_prompt = """You are PlantMind AI, summarizing an internal IT document for a
before/after comparison view an employee and admin will read side by side.
Reply with ONLY a compact JSON object, nothing else - no markdown fences, no commentary:
{"summary": "<2-3 sentence plain-text summary>", "key_points": ["<point 1>", "<point 2>", "<point 3>"]}
Keep key_points to at most 5 short, concrete bullet points (procedures, settings,
version numbers, deadlines, responsibilities, etc.) - skip it (empty list) if the
document is too short or generic to have distinct key points."""

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Document name: {doc_name}\n\nContent:\n{text[:6000]}"},
            ],
            temperature=0.2,
            max_tokens=400,
        )
        raw = (completion.choices[0].message.content or "").strip()
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        parsed = json.loads(raw)
        summary = str(parsed.get("summary", "")).strip() or "No summary available."
        key_points = parsed.get("key_points", [])
        if not isinstance(key_points, list):
            key_points = []
        key_points = [str(p).strip() for p in key_points if str(p).strip()][:5]

        return {"summary": summary, "key_points": key_points}
    except Exception as e:
        print("[WARN] analyze_text_with_ai failed:", e)
        return {
            "summary": f"PlantMind AI could not analyze this file automatically ({str(e)}).",
            "key_points": [],
        }


# --------------------------------------------------------------------------
# AI UPLOAD SCREENING  ->  used for REPLACEMENT uploads of an existing
# document (propose_document_update). Also feeds the Decision Engine.
# --------------------------------------------------------------------------
def score_upload_with_ai(filename, content_text, doc_name):
    """
    Ask Groq to evaluate a freshly uploaded replacement file against
    OnlyTech's knowledge-base standards and return a 0-100 readiness score
    plus the Decision Engine's routing decision.

    Returns: {"decision": "approve" | "review" | "reject", "score": int, "reasoning": str}

    Falls back to "review" (never auto-reject/auto-approve) whenever the AI
    genuinely can't evaluate the file (Groq unavailable, no extractable
    text, request error) - those are infrastructure gaps, not a real
    low-quality signal, so a human should still take a look.
    """
    if groq_client is None:
        return {"decision": "review", "score": 0,
                "reasoning": "PlantMind AI is not configured on the server, so this needs a manual look."}

    text_for_ai = (content_text or "").strip()
    if not text_for_ai:
        return {"decision": "review", "score": 0,
                "reasoning": "No readable text could be extracted from this file (likely an image, spreadsheet, or scanned document)."}

    system_prompt = """You are a compliance reviewer for OnlyTech's internal IT knowledge base.
Score how ready a submitted document is to be published as-is, based on:
- Is it a real, coherent IT/operations document (SOP, policy, runbook, guide, report, etc.)?
- Is it free of placeholder, junk, spam, or test content?
- Does it look complete, professional, and safe to publish without edits?
- Since this is a REPLACEMENT for an existing document, does it stay on-topic and consistent
  with the kind of document it is replacing (not a completely unrelated file)?

Reply with ONLY a compact JSON object, nothing else - no markdown fences, no commentary:
{"score": <integer 0-100>, "reasoning": "<one or two plain sentences>"}"""

    try:
        completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Proposed document name: {doc_name}\nOriginal filename: {filename}\n\nExtracted content:\n{text_for_ai[:6000]}"},
            ],
            temperature=0.1,
            max_tokens=250,
        )
        raw = (completion.choices[0].message.content or "").strip()
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        parsed = json.loads(raw)
        score = int(parsed.get("score", 50))
        score = max(0, min(100, score))
        reasoning = str(parsed.get("reasoning", "")).strip()[:500]

        if not reasoning:
            reasoning = "PlantMind AI evaluated this document automatically."

        return {"decision": decide_from_score(score), "score": score, "reasoning": reasoning}
    except Exception as e:
        print("[WARN] AI upload scoring failed:", e)
        return {"decision": "review", "score": 0,
                "reasoning": f"PlantMind AI could not score this file automatically ({str(e)})."}


# --------------------------------------------------------------------------
# HEALTH CHECK
# --------------------------------------------------------------------------
@app.route("/api/health", methods=["GET"])
def health():
    return jsonify({
        "status": "ok",
        "groq_configured": groq_client is not None,
        "groq_model": GROQ_MODEL_NAME,
        "supabase_configured": supabase is not None,
        "storage_bucket": STORAGE_BUCKET,
        "pypdf_available": PYPDF_AVAILABLE,
        "docx_available": DOCX_AVAILABLE,
        "groq_comparison_enabled": groq_client is not None,
        "ai_auto_approve_min_score": AI_AUTO_APPROVE_MIN_SCORE,
        "ai_auto_reject_max_score": AI_AUTO_REJECT_MAX_SCORE,
        "auto_decision_routing_enabled": True,
        "time": now_iso()
    })


# --------------------------------------------------------------------------
# AUTH  ->  table: users
# --------------------------------------------------------------------------
@app.route("/api/auth/register", methods=["POST"])
@require_supabase
def register():
    data = request.get_json(force=True) or {}
    required = ["employee_id", "name", "email", "department", "role", "password"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    if not data["email"].lower().endswith("@onlytech.com"):
        return jsonify({"error": "Registration is restricted to @onlytech.com emails"}), 400

    existing = supabase.table("users").select("employee_id").eq("employee_id", data["employee_id"]).execute()
    if existing.data:
        return jsonify({"error": "Employee ID already registered"}), 409

    record = {
        "employee_id": data["employee_id"],
        "name": data["name"],
        "email": data["email"],
        "department": data["department"],
        "designation": data.get("designation", data["role"]),
        "role": data["role"],
        "password_hash": generate_password_hash(data["password"]),
        "twofa_enabled": False,
        "created_at": now_iso(),
    }
    supabase.table("users").insert(record).execute()
    record.pop("password_hash")
    return jsonify({"message": "Account created", "user": record}), 201


@app.route("/api/auth/login", methods=["POST"])
@require_supabase
def login():
    data = request.get_json(force=True) or {}
    employee_id = data.get("employee_id")
    password    = data.get("password")
    if not employee_id or not password:
        return jsonify({"error": "employee_id and password are required"}), 400

    if employee_id == "ADMIN_01" and password == "AdminPass123!":
        res = supabase.table("users").select("*").eq("employee_id", "ADMIN_01").execute()
        if not res.data:
            admin_record = {
                "employee_id": "ADMIN_01",
                "name": "Super Admin",
                "email": "admin@onlytech.com",
                "department": "IT Operations",
                "designation": "Systems Architect",
                "role": "Admin",
                "password_hash": generate_password_hash("AdminPass123!"),
                "twofa_enabled": False,
                "created_at": now_iso(),
            }
            supabase.table("users").insert(admin_record).execute()
            admin_record.pop("password_hash", None)
            return jsonify({"message": "Local Admin Session Created & Seeded", "user": admin_record}), 200

    res = supabase.table("users").select("*").eq("employee_id", employee_id).execute()
    if not res.data:
        return jsonify({"error": "Invalid Employee ID or password"}), 401

    user = res.data[0]
    if not check_password_hash(user["password_hash"], password):
        return jsonify({"error": "Invalid Employee ID or password"}), 401

    user.pop("password_hash", None)
    return jsonify({"message": "Login successful", "user": user}), 200


# --------------------------------------------------------------------------
# DOCUMENTS  ->  tables: documents, document_history
# --------------------------------------------------------------------------
@app.route("/api/documents", methods=["GET"])
@require_supabase
def list_documents():
    uploaded_by = (request.args.get("uploaded_by") or "").strip()
    q = supabase.table("documents").select("*")
    if uploaded_by:
        q = q.eq("uploaded_by", uploaded_by)
    res = q.order("updated_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/documents/<doc_id>", methods=["GET"])
@require_supabase
def get_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    history = supabase.table("document_history").select("*").eq("document_id", doc_id).order("created_at", desc=True).execute()
    result = doc.data[0]
    result["history"] = history.data
    return jsonify(result), 200


@app.route("/api/documents", methods=["POST"])
@require_supabase
def create_document():
    """JSON-based creation. Used by admin tools / scripts, not the employee upload flow."""
    data = request.get_json(force=True) or {}
    if not data.get("name"):
        return jsonify({"error": "Document name is required"}), 400

    doc_id = str(uuid.uuid4())
    record = {
        "id": doc_id,
        "name": data["name"],
        "icon": pick_icon(data["name"]),
        "version": data.get("version", "v1.0"),
        "status": data.get("status", "Active"),
        "summary": data.get("summary", ""),
        "content": data.get("content", ""),
        "uploaded_by": data.get("uploaded_by", "Unknown"),
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    supabase.table("documents").insert(record).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "version": record["version"],
        "updated_by": record["uploaded_by"],
        "changes": "Initial upload",
        "content_snapshot": record["content"],
        "created_at": now_iso(),
    }).execute()
    log_activity(record["uploaded_by"], "uploaded_document", record["name"])
    return jsonify(record), 201


@app.route("/api/documents/summarize", methods=["POST"])
def summarize_document_preview():
    """
    PREVIEW-ONLY endpoint for the plain "Upload File" button in Documents.
    Extracts text (plain text, PDF, or DOCX) and asks Groq for a short
    2-3 sentence summary. Nothing is saved to Supabase here.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)

    if not content_text.strip():
        size_kb = len(file_bytes) / 1024
        return jsonify({
            "summary": f"\"{file.filename}\" is a {mimetype or 'binary'} file (~{size_kb:.1f} KB). "
                       f"Automatic text summarization isn't available for this file type, "
                       f"but the file itself will still be submitted for review.",
            "content_preview": "",
            "can_summarize": False
        }), 200

    result = analyze_text_with_ai(file.filename, content_text)
    return jsonify({
        "summary": result["summary"],
        "content_preview": content_text[:500],
        "can_summarize": True
    }), 200


@app.route("/api/documents/analyze-file", methods=["POST"])
def analyze_uploaded_file():
    """
    PREVIEW-ONLY endpoint used by the "Update Docs" dashboard's right-hand
    ("Proposed New Version") panel. Given a replacement file (multipart,
    field 'file', plus optional 'doc_name'), extracts its text and returns
    an AI summary + key points WITHOUT saving anything. The frontend calls
    /api/documents/<id>/propose-update separately once the employee confirms.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    doc_name = request.form.get("doc_name") or file.filename

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    text = extract_text(file_bytes, mimetype, file.filename)

    result = analyze_text_with_ai(doc_name, text)
    return jsonify(result), 200


@app.route("/api/documents/<doc_id>/analyze", methods=["GET"])
@require_supabase
def analyze_current_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404

    d = doc.data[0]
    text = ensure_doc_content(d)
    result = analyze_text_with_ai(d.get("name", "Document"), text or d.get("summary", ""))
    return jsonify(result), 200


@app.route("/api/documents/propose-upload", methods=["POST"])
@require_supabase
def propose_upload():
    """
    Employee-facing NEW document upload endpoint (plain Documents tab).

    Runs the full Decision Engine pipeline:
      1. Extract text from the uploaded file
      2. Search the Knowledge Base for the best TITLE match
      3. Groq-COMPARE the uploaded content against that matched document's content
      4. Combine into a single match score
      5. Route: >= approve threshold -> auto-publish immediately
                <  reject threshold  -> auto-reject immediately (file deleted)
                in between           -> Pending row in the admin Approvals tab
      6. Every outcome is written to document_approvals + activity_log so the
         admin dashboard (Overview stats, Recent Activity, Approvals tab)
         updates automatically without a manual refresh.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    uploaded_by = request.form.get("uploaded_by", "Unknown")
    name = request.form.get("name") or file.filename
    summary = request.form.get("summary", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename, folder="approvals")

    try:
        ai_result = score_upload_by_title_match(file.filename, name, content_text or summary)
    except Exception as e:
        ai_result = {"decision": "review", "score": 0,
                     "reasoning": f"Title-match scoring failed: {e}", "matched_document": None}

    matched_doc = ai_result.get("matched_document")
    decision = ai_result.get("decision", "review")

    if matched_doc:
        old_version = matched_doc.get("version", "v1.0")
        try:
            num = float(old_version.replace("v", ""))
            proposed_version = f"v{round(num + 0.1, 1)}"
        except Exception:
            proposed_version = "v2.0"
        old_content = ensure_doc_content(matched_doc)
        document_id = matched_doc["id"]
    else:
        old_version = "v0.0"
        proposed_version = "v1.0"
        old_content = ""
        document_id = None

    pending_record = {
        "id": str(uuid.uuid4()),
        "document_id": document_id,
        "doc_name": name,
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "submitted_by": uploaded_by,
        "status": "Pending",
        "created_at": now_iso(),
    }

    if decision == "approve":
        publish_pending_record(pending_record, AUTO_APPROVER_NAME)
        return jsonify({
            "outcome": "approved",
            "message": f'"{name}" matched "{matched_doc["name"]}" at {ai_result["score"]}% and was '
                       f'automatically published as {proposed_version}.',
            "ai": ai_result,
        }), 200

    if decision == "reject":
        discard_pending_record(pending_record, AUTO_REJECTOR_NAME, ai_result["reasoning"])
        return jsonify({
            "outcome": "rejected",
            "message": f'"{name}" scored only {ai_result["score"]}% against the Knowledge Base and was '
                       f'automatically rejected. {ai_result["reasoning"]}',
            "ai": ai_result,
        }), 200

    # decision == "review"
    insert_approval_audit_row(pending_record, ai_result)
    log_activity(
        uploaded_by, "proposed_edit", name,
        f"Submitted for admin review — match score {ai_result['score']}%: {ai_result['reasoning']}"
    )
    return jsonify({
        "outcome": "review",
        "message": f'"{name}" scored {ai_result["score"]}% (needs a human look) and is now pending admin approval.',
        "ai": ai_result,
    }), 202


@app.route("/api/documents/<doc_id>/propose-update", methods=["POST"])
@require_supabase
def propose_document_update(doc_id):
    """
    Employee-facing REPLACEMENT upload for an EXISTING document, from the
    "Update Docs" dashboard.

    Runs the same Decision Engine as propose_upload, but the score comes
    from PlantMind AI's content-quality/consistency read (since the target
    document is already known - there's no title match to do).
    """
    existing = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not existing.data:
        return jsonify({"error": "Document not found"}), 404
    current = existing.data[0]

    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    submitted_by = request.form.get("submitted_by", "Unknown")
    reason = request.form.get("reason", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    new_content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename, folder="approvals")

    old_version = current.get("version", "v1.0")
    try:
        num = float(old_version.replace("v", ""))
        proposed_version = f"v{round(num + 0.1, 1)}"
    except Exception:
        proposed_version = "v2.0"

    old_content = ensure_doc_content(current)

    try:
        ai_result = score_upload_with_ai(file.filename, new_content_text, current["name"])
    except Exception as e:
        ai_result = {"decision": "review", "score": 0, "reasoning": f"AI scoring failed: {e}"}

    decision = ai_result.get("decision", "review")

    pending_record = {
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "doc_name": current["name"],
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": new_content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "submitted_by": submitted_by,
        "status": "Pending",
        "created_at": now_iso(),
    }

    if decision == "approve":
        publish_pending_record(pending_record, AUTO_APPROVER_NAME)
        return jsonify({
            "outcome": "approved",
            "message": f"PlantMind AI scored this {ai_result['score']}/100 and automatically published it as {proposed_version}.",
            "approval_id": pending_record["id"], "proposed_version": proposed_version, "ai": ai_result,
        }), 200

    if decision == "reject":
        discard_pending_record(pending_record, AUTO_REJECTOR_NAME, reason or ai_result["reasoning"])
        return jsonify({
            "outcome": "rejected",
            "message": f"PlantMind AI scored this only {ai_result['score']}/100 and it was automatically rejected. {ai_result['reasoning']}",
            "approval_id": pending_record["id"], "ai": ai_result,
        }), 200

    # decision == "review"
    insert_approval_audit_row(pending_record, ai_result)
    log_activity(submitted_by, "proposed_edit", current["name"],
                 f"Needs admin review — PlantMind AI score {ai_result['score']}/100: {reason or ai_result['reasoning']}")
    return jsonify({
        "outcome": "review",
        "message": f"Update submitted — PlantMind AI scored this {ai_result['score']}/100. An admin will review it before it's published.",
        "approval_id": pending_record["id"], "proposed_version": proposed_version, "ai": ai_result,
    }), 201


@app.route("/api/documents/upload", methods=["POST"])
@require_supabase
def upload_document():
    """
    Direct admin upload -> stored in Supabase Storage, text extracted, published
    immediately with no approval step. Used by the admin Command Center's
    Knowledge Base "Upload File" button ONLY (admin is already the approver).
    Employee uploads go through /api/documents/propose-upload or
    /api/documents/<id>/propose-update instead, which run the Decision Engine.
    """
    if "file" not in request.files:
        return jsonify({"error": "file is required"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    uploaded_by = request.form.get("uploaded_by", "Super Admin")
    name = request.form.get("name") or file.filename
    summary = request.form.get("summary", "")

    file_bytes = file.read()
    mimetype = file.mimetype or "application/octet-stream"
    content_text = extract_text(file_bytes, mimetype, file.filename)
    unique_path = upload_bytes_to_storage(file_bytes, mimetype, file.filename)

    doc_id = str(uuid.uuid4())
    record = {
        "id": doc_id,
        "name": name,
        "icon": pick_icon(file.filename),
        "version": "v1.0",
        "status": "Active",
        "summary": summary,
        "content": content_text,
        "file_path": unique_path,
        "file_name": file.filename,
        "file_type": mimetype,
        "file_size": len(file_bytes),
        "uploaded_by": uploaded_by,
        "created_at": now_iso(),
        "updated_at": now_iso(),
    }
    supabase.table("documents").insert(record).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()), "document_id": doc_id, "version": "v1.0",
        "updated_by": uploaded_by, "changes": "Initial upload",
        "content_snapshot": content_text, "created_at": now_iso(),
    }).execute()
    log_activity(uploaded_by, "uploaded_document", name, f"Uploaded {file.filename}")
    return jsonify(record), 201


@app.route("/api/documents/<doc_id>/download", methods=["GET"])
@require_supabase
def download_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    d = doc.data[0]
    if not d.get("file_path"):
        return jsonify({"error": "No file stored for this document"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(d["file_path"])
    return Response(
        file_bytes,
        mimetype=d.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{d.get("file_name") or d["name"]}"'}
    )


@app.route("/api/documents/<doc_id>", methods=["DELETE"])
@require_supabase
def delete_document(doc_id):
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404
    d = doc.data[0]

    delete_from_storage(d.get("file_path"))

    try:
        supabase.table("document_history").delete().eq("document_id", doc_id).execute()
    except Exception as e:
        print("[WARN] Could not delete document history:", e)

    supabase.table("documents").delete().eq("id", doc_id).execute()

    deleted_by = request.args.get("deleted_by")
    if not deleted_by:
        deleted_by = (request.get_json(silent=True) or {}).get("deleted_by", "Super Admin")
    log_activity(deleted_by, "deleted_document", d["name"], "Document permanently removed from knowledge base")

    return jsonify({"message": "Document deleted", "id": doc_id}), 200


@app.route("/api/documents/<doc_id>/version", methods=["POST"])
@require_supabase
def update_document_version(doc_id):
    """Direct admin version bump."""
    data = request.get_json(force=True) or {}
    doc = supabase.table("documents").select("*").eq("id", doc_id).execute()
    if not doc.data:
        return jsonify({"error": "Document not found"}), 404

    current = doc.data[0]
    try:
        num = float(current["version"].replace("v", ""))
        new_version = f"v{round(num + 0.1, 1)}"
    except Exception:
        new_version = "v2.0"

    updates = {"version": new_version, "status": "Updated", "updated_at": now_iso()}
    if data.get("content"):
        updates["content"] = data["content"]

    supabase.table("documents").update(updates).eq("id", doc_id).execute()
    supabase.table("document_history").insert({
        "id": str(uuid.uuid4()), "document_id": doc_id, "version": new_version,
        "updated_by": data.get("updated_by", "Unknown"),
        "changes": data.get("changes", "Updated document"),
        "content_snapshot": data.get("content", ""), "created_at": now_iso(),
    }).execute()
    log_activity(data.get("updated_by", "Unknown"), "uploaded_document", current["name"], "Version updated")
    return jsonify({"message": "Version updated", "version": new_version}), 200


@app.route("/api/documents/backfill-content", methods=["POST"])
@require_supabase
def backfill_document_content():
    docs = supabase.table("documents").select("*").execute().data or []
    updated, skipped, failed = [], [], []

    for doc in docs:
        if (doc.get("content") or "").strip():
            skipped.append(doc["name"])
            continue
        if not doc.get("file_path"):
            skipped.append(doc["name"])
            continue
        text = ensure_doc_content(doc)
        if text:
            updated.append(doc["name"])
        else:
            failed.append(doc["name"])

    return jsonify({
        "updated": updated,
        "skipped_already_had_content_or_no_file": skipped,
        "failed_to_extract": failed,
    }), 200


# --------------------------------------------------------------------------
# DOCUMENT APPROVALS  ->  table: document_approvals
# --------------------------------------------------------------------------
@app.route("/api/documents/propose", methods=["POST"])
@require_supabase
def propose_document():
    """Propose an edit to an existing document (multipart/form-data)."""
    submitted_by = request.form.get("submitted_by", "Unknown")
    doc_id = request.form.get("document_id") or None
    doc_name = (request.form.get("name") or "").strip()
    new_content_text = request.form.get("content", "")

    file = request.files.get("file")
    file_path = file_name = file_type = None
    file_size = None
    if file and file.filename:
        file_bytes = file.read()
        file_type = file.mimetype or "application/octet-stream"
        file_name = file.filename
        file_size = len(file_bytes)
        file_path = upload_bytes_to_storage(file_bytes, file_type, file.filename, folder="approvals")
        if not new_content_text:
            new_content_text = extract_text(file_bytes, file_type, file.filename)

    old_content = ""
    old_version = "v0.0"
    if doc_id:
        existing = supabase.table("documents").select("*").eq("id", doc_id).execute()
        if not existing.data:
            return jsonify({"error": "Document not found"}), 404
        old_content = existing.data[0].get("content", "") or ""
        old_version = existing.data[0].get("version", "v1.0")
        doc_name = doc_name or existing.data[0]["name"]
        try:
            num = float(old_version.replace("v", ""))
            proposed_version = f"v{round(num + 0.1, 1)}"
        except Exception:
            proposed_version = "v2.0"
    else:
        proposed_version = "v1.0"

    if not doc_name:
        return jsonify({"error": "Document name is required"}), 400

    approval_id = str(uuid.uuid4())
    record = {
        "id": approval_id,
        "document_id": doc_id,
        "doc_name": doc_name,
        "old_version": old_version,
        "proposed_version": proposed_version,
        "old_content": old_content,
        "new_content": new_content_text,
        "file_path": file_path,
        "file_name": file_name,
        "file_type": file_type,
        "file_size": file_size,
        "submitted_by": submitted_by,
        "status": "Pending",
        "created_at": now_iso(),
    }
    supabase.table("document_approvals").insert(record).execute()
    log_activity(submitted_by, "proposed_edit", doc_name, f"Proposed {proposed_version}")
    return jsonify(record), 201


@app.route("/api/approvals", methods=["POST"])
@require_supabase
def propose_edit_json():
    """
    JSON-based edit proposal, used by the employee portal's "Propose New
    Version" dialog on the Documents tab (posts JSON, not multipart).
    """
    data = request.get_json(force=True) or {}
    doc_id = data.get("doc_id")
    doc_name = (data.get("doc_name") or "").strip()
    if not doc_id or not doc_name:
        return jsonify({"error": "doc_id and doc_name are required"}), 400

    record = {
        "id": str(uuid.uuid4()),
        "document_id": doc_id,
        "doc_name": doc_name,
        "old_version": data.get("old_version", "v1.0"),
        "proposed_version": data.get("proposed_version", "v1.1"),
        "old_content": data.get("old_content", ""),
        "new_content": data.get("new_content", ""),
        "file_path": None,
        "file_name": data.get("file_name"),
        "file_size": data.get("file_size"),
        "file_type": None,
        "submitted_by": data.get("submitted_by", "Unknown"),
        "status": "Pending",
        "created_at": now_iso(),
    }
    supabase.table("document_approvals").insert(record).execute()
    log_activity(record["submitted_by"], "proposed_edit", doc_name, data.get("changes", "Proposed an edit"))
    return jsonify(record), 201


@app.route("/api/approvals", methods=["GET"])
@require_supabase
def list_approvals():
    status = request.args.get("status")
    submitted_by = request.args.get("submitted_by")

    q = supabase.table("document_approvals").select("*")
    if status:
        q = q.eq("status", status)
    if submitted_by:
        q = q.eq("submitted_by", submitted_by)
    res = q.order("created_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/approvals/<approval_id>", methods=["GET"])
@require_supabase
def get_approval(approval_id):
    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Approval not found"}), 404
    return jsonify(res.data[0]), 200


@app.route("/api/approvals/<approval_id>/file", methods=["GET"])
@require_supabase
def download_approval_file(approval_id):
    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Not found"}), 404
    a = res.data[0]
    if not a.get("file_path"):
        return jsonify({"error": "No file attached"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(a["file_path"])
    return Response(
        file_bytes,
        mimetype=a.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{a.get("file_name") or "file"}"'}
    )


@app.route("/api/approvals/<approval_id>/decide", methods=["POST"])
@require_supabase
def decide_approval(approval_id):
    """
    Manual admin decision for whatever landed in the review band (70-89%).
    Approving publishes the file to the Knowledge Base (updating the
    matched document if document_id is set, otherwise creating a brand new
    document); rejecting deletes the uploaded file and removes the request.
    """
    data = request.get_json(force=True) or {}
    approve = bool(data.get("approve"))
    decided_by = data.get("decided_by", "Super Admin")

    res = supabase.table("document_approvals").select("*").eq("id", approval_id).execute()
    if not res.data:
        return jsonify({"error": "Approval not found"}), 404
    app_row = res.data[0]

    if app_row["status"] != "Pending":
        return jsonify({"error": "This request was already decided"}), 409

    if approve:
        if app_row["document_id"]:
            doc_id = app_row["document_id"]
            updates = {
                "version": app_row["proposed_version"],
                "status": "Active",
                "content": app_row["new_content"],
                "updated_at": now_iso(),
            }
            if app_row.get("file_path"):
                updates.update({
                    "file_path": app_row["file_path"],
                    "file_name": app_row["file_name"],
                    "file_type": app_row["file_type"],
                    "file_size": app_row["file_size"],
                })
            supabase.table("documents").update(updates).eq("id", doc_id).execute()
        else:
            doc_id = str(uuid.uuid4())
            supabase.table("documents").insert({
                "id": doc_id,
                "name": app_row["doc_name"],
                "icon": pick_icon(app_row.get("file_name") or app_row["doc_name"]),
                "version": app_row["proposed_version"],
                "status": "Active",
                "summary": "",
                "content": app_row["new_content"],
                "file_path": app_row.get("file_path"),
                "file_name": app_row.get("file_name"),
                "file_type": app_row.get("file_type"),
                "file_size": app_row.get("file_size"),
                "uploaded_by": app_row["submitted_by"],
                "created_at": now_iso(),
                "updated_at": now_iso(),
            }).execute()

        supabase.table("document_history").insert({
            "id": str(uuid.uuid4()), "document_id": doc_id, "version": app_row["proposed_version"],
            "updated_by": app_row["submitted_by"], "changes": "Approved edit request.",
            "content_snapshot": app_row["new_content"], "created_at": now_iso(),
        }).execute()

        supabase.table("document_approvals").update({
            "status": "Approved", "decided_at": now_iso(), "decided_by": decided_by,
        }).eq("id", approval_id).execute()

        log_activity(decided_by, "approved_edit", app_row["doc_name"], f"Published {app_row['proposed_version']}")
    else:
        if app_row.get("file_path"):
            delete_from_storage(app_row["file_path"])

        supabase.table("document_approvals").delete().eq("id", approval_id).execute()
        log_activity(decided_by, "rejected_edit", app_row["doc_name"], "Edit request rejected and file removed")

    return jsonify({"message": "Decision recorded", "approve": approve}), 200


@app.route("/api/approvals/automate-batch", methods=["POST"])
@require_supabase
def automate_batch_approvals():
    """
    Automate processing of ALL currently-Pending approvals using the
    Groq-powered Decision Engine (score_approval_with_groq):

    1. For each pending row, resolve the Knowledge Base document it's about
       (via document_id, or a fresh title match for brand-new uploads).
    2. Pull that document's CURRENT text and the approval's PROPOSED text
       (extracting from Storage on demand if needed).
    3. Ask Groq to score how consistent the two are (0-100), combine with a
       title-match score.
    4. Decision Engine:
         Score >= AI_AUTO_APPROVE_MIN_SCORE  -> Auto Publish
         Score in the review band            -> stays Pending (needs a human)
         Score <  AI_AUTO_REJECT_MAX_SCORE   -> Auto Reject (file deleted)
    5. Every outcome is written back to document_approvals + activity_log,
       so Overview / Recent Activity / the Approvals tab update automatically.

    Returns: {"processed", "approved", "rejected", "reviewed", "results": [...]}
    """
    data = request.get_json(force=True) or {}
    decided_by = data.get("decided_by", "PlantMind AI (Batch Auto)")

    pending = supabase.table("document_approvals").select("*").eq("status", "Pending") \
        .order("created_at", desc=False).execute().data or []

    if not pending:
        return jsonify({
            "processed": 0, "approved": 0, "rejected": 0, "reviewed": 0,
            "results": [], "message": "No pending approvals to process"
        }), 200

    results = []
    approved_count = 0
    rejected_count = 0
    reviewed_count = 0

    for approval in pending:
        doc_name = approval.get("doc_name", "Unknown")
        try:
            ai_result = score_approval_with_groq(approval)
            score = ai_result.get("score", 0)
            reasoning = ai_result.get("reasoning", "Batch automated scoring")
            matched_doc = ai_result.get("matched_document")

            decision = decide_from_score(score)
            # Never auto-approve a brand-new upload that has nothing in the
            # Knowledge Base to actually update.
            if decision == "approve" and not matched_doc:
                decision = "review"

            if decision == "approve":
                publish_pending_record(approval, decided_by)
                approved_count += 1
                results.append({"id": approval["id"], "doc_name": doc_name, "decision": "approved",
                                 "score": score, "reason": reasoning})

            elif decision == "reject":
                discard_pending_record(approval, decided_by, reasoning)
                rejected_count += 1
                results.append({"id": approval["id"], "doc_name": doc_name, "decision": "rejected",
                                 "score": score, "reason": reasoning})

            else:  # review — leave it Pending, just record the latest score for the admin UI
                reviewed_count += 1
                try:
                    supabase.table("document_approvals").update({
                        "ai_score": score, "ai_reasoning": reasoning,
                    }).eq("id", approval["id"]).execute()
                except Exception:
                    pass  # columns may not exist on an older schema — non-fatal
                results.append({"id": approval["id"], "doc_name": doc_name, "decision": "review",
                                 "score": score, "reason": reasoning})

        except Exception as e:
            print(f"[WARN] Batch processing failed for {approval.get('id')}: {e}")
            results.append({"id": approval.get("id"), "doc_name": doc_name, "decision": "error", "reason": str(e)})

    log_activity(
        decided_by, "batch_automate_approvals", "Batch Processing",
        f"Processed {len(pending)} pending approvals: {approved_count} approved, "
        f"{rejected_count} rejected, {reviewed_count} kept for review"
    )

    return jsonify({
        "processed": len(pending),
        "approved": approved_count,
        "rejected": rejected_count,
        "reviewed": reviewed_count,
        "results": results,
        "message": f"Batch automation complete: {approved_count} approved, "
                   f"{rejected_count} rejected, {reviewed_count} need human review"
    }), 200


# --------------------------------------------------------------------------
# MAINTENANCE TICKETS  ->  table: maintenance_tickets
# --------------------------------------------------------------------------
@app.route("/api/maintenance", methods=["GET"])
@require_supabase
def list_tickets():
    res = supabase.table("maintenance_tickets").select("*").order("created_at", desc=True).execute()
    return jsonify(res.data), 200


@app.route("/api/maintenance", methods=["POST"])
@require_supabase
def create_ticket():
    is_multipart = bool(request.content_type) and "multipart/form-data" in request.content_type
    data = request.form if is_multipart else (request.get_json(force=True) or {})

    required = ["system_name", "issue_type", "description", "severity", "reported_by"]
    missing = [f for f in required if not data.get(f)]
    if missing:
        return jsonify({"error": f"Missing fields: {', '.join(missing)}"}), 400

    file_path = file_name = file_type = None
    file_size = None
    if is_multipart:
        file = request.files.get("file")
        if file and file.filename:
            file_bytes = file.read()
            file_type = file.mimetype or "application/octet-stream"
            file_name = file.filename
            file_size = len(file_bytes)
            file_path = upload_bytes_to_storage(file_bytes, file_type, file.filename, folder="tickets")

    count_res = supabase.table("maintenance_tickets").select("id", count="exact").execute()
    next_num = (count_res.count or 0) + 1042

    record = {
        "id": str(uuid.uuid4()),
        "ticket_id": f"INC-{next_num}",
        "system_name": data["system_name"],
        "issue_type": data["issue_type"],
        "severity": data["severity"],
        "description": data["description"],
        "reported_by": data["reported_by"],
        "status": "Open",
        "resolution_notes": None,
        "file_path": file_path,
        "file_name": file_name,
        "file_type": file_type,
        "file_size": file_size,
        "created_at": now_iso(),
        "resolved_at": None,
    }
    supabase.table("maintenance_tickets").insert(record).execute()
    log_activity(data["reported_by"], "reported_incident", record["ticket_id"], data["system_name"])
    return jsonify(record), 201


@app.route("/api/maintenance/<ticket_id>/file", methods=["GET"])
@require_supabase
def download_ticket_file(ticket_id):
    ticket = supabase.table("maintenance_tickets").select("*").eq("id", ticket_id).execute()
    if not ticket.data:
        return jsonify({"error": "Ticket not found"}), 404
    t = ticket.data[0]
    if not t.get("file_path"):
        return jsonify({"error": "No file attached to this ticket"}), 404
    file_bytes = supabase.storage.from_(STORAGE_BUCKET).download(t["file_path"])
    return Response(
        file_bytes,
        mimetype=t.get("file_type") or "application/octet-stream",
        headers={"Content-Disposition": f'inline; filename="{t.get("file_name") or "attachment"}"'}
    )


@app.route("/api/maintenance/<ticket_id>/resolve", methods=["POST"])
@require_supabase
def resolve_ticket(ticket_id):
    data = request.get_json(force=True) or {}
    notes = data.get("resolution_notes")
    if not notes:
        return jsonify({"error": "resolution_notes is required"}), 400

    ticket = supabase.table("maintenance_tickets").select("*").eq("id", ticket_id).execute()
    if not ticket.data:
        return jsonify({"error": "Ticket not found"}), 404

    supabase.table("maintenance_tickets").update({
        "status": "Resolved",
        "resolution_notes": notes,
        "resolved_at": now_iso(),
        "resolved_by": data.get("resolved_by", "Admin"),
    }).eq("id", ticket_id).execute()

    log_activity(data.get("resolved_by", "Admin"), "resolved_incident", ticket.data[0]["ticket_id"], notes)
    return jsonify({"message": "Ticket resolved"}), 200


# --------------------------------------------------------------------------
# ACTIVITY LOG  ->  table: activity_log  (drives Recent Activity Feed)
# --------------------------------------------------------------------------
@app.route("/api/activity", methods=["GET"])
@require_supabase
def list_activity():
    limit = int(request.args.get("limit", 10))
    res = supabase.table("activity_log").select("*").order("created_at", desc=True).limit(limit).execute()
    return jsonify(res.data), 200


# --------------------------------------------------------------------------
# DASHBOARD STATS  ->  computed from documents / approvals / tickets / ai logs
# --------------------------------------------------------------------------
@app.route("/api/stats/overview", methods=["GET"])
@require_supabase
def stats_overview():
    docs = supabase.table("documents").select("id,status").execute().data or []
    approvals = supabase.table("document_approvals").select("id").eq("status", "Pending").execute().data or []
    tickets = supabase.table("maintenance_tickets").select("id,status,severity,created_at").execute().data or []
    ai_logs = supabase.table("ai_query_logs").select("sources,created_at").order("created_at", desc=True).limit(200).execute().data or []

    active_documents = len([d for d in docs if d.get("status") == "Active"])
    pending_approvals = len(approvals)
    open_incidents = len([t for t in tickets if t.get("status") != "Resolved"])
    critical_open = len([t for t in tickets if t.get("status") != "Resolved" and "P1" in (t.get("severity") or "")])

    today = datetime.datetime.utcnow().date()
    days = [today - datetime.timedelta(days=i) for i in range(6, -1, -1)]
    trend = []
    for d in days:
        count = 0
        for t in tickets:
            try:
                created = datetime.datetime.fromisoformat(t["created_at"]).date()
            except Exception:
                continue
            if created == d:
                count += 1
        trend.append({"day": d.strftime("%a")[0], "count": count})

    counter = collections.Counter()
    for log in ai_logs:
        for s in (log.get("sources") or "").split(", "):
            s = s.strip()
            if s:
                counter[s] += 1
    total = sum(counter.values()) or 1
    top_queries = [
        {"name": name, "pct": round(count / total * 100)}
        for name, count in counter.most_common(3)
    ]

    return jsonify({
        "active_documents": active_documents,
        "pending_approvals": pending_approvals,
        "open_incidents": open_incidents,
        "system_alerts": critical_open,
        "incident_trend": trend,
        "top_queries": top_queries,
    }), 200


# --------------------------------------------------------------------------
# PLANTMIND AI  ->  Groq, grounded on real extracted document text +
# approvals + version history + activity log + tickets + employees
# --------------------------------------------------------------------------
MODE_INSTRUCTIONS = {
    "general":    "Answer the question clearly and concisely for an employee using the company knowledge base.",
    "root_cause": "Perform a root cause analysis. Identify the most likely technical cause(s) of the issue described.",
    "policy":     "Answer strictly from a compliance/policy point of view, quoting relevant rules where applicable.",
    "summary":    "Give a short executive summary in 2-4 bullet points.",
}


def find_relevant_documents(question, limit=3):
    """
    Search the documents table for relevant records. NOTE: this intentionally
    no longer filters to status == "Active" - Draft/Updated documents are
    included too, so PlantMind AI can talk about ANY document in the
    Knowledge Base, not just the currently-published one.
    """
    if supabase is None:
        return []
    try:
        res = supabase.table("documents") \
            .select("id,name,content,summary,version,status,file_path,file_type,file_name") \
            .execute()
        docs = res.data or []
    except Exception:
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)

    scored = []
    for doc in docs:
        text = f"{doc.get('name','')} {doc.get('summary','')} {doc.get('content','')}".lower()
        score = sum(1 for w in q_words if w in text)
        if score == 0 and not (doc.get("content") or "").strip() and doc.get("file_path"):
            name_summary = f"{doc.get('name','')} {doc.get('summary','')}".lower()
            score = sum(1 for w in q_words if w in name_summary)
        if score > 0:
            scored.append((score, doc))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [d for _, d in scored[:limit]]


def find_relevant_incidents(question, limit=3):
    if supabase is None:
        return []
    try:
        res = supabase.table("maintenance_tickets") \
            .select("id,ticket_id,system_name,issue_type,severity,description,status,resolution_notes,reported_by,resolved_by,created_at,resolved_at") \
            .execute()
        tickets = res.data or []
    except Exception:
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)
    if not q_words:
        return []

    scored = []
    for ticket in tickets:
        text = " ".join([
            ticket.get("ticket_id") or "",
            ticket.get("system_name") or "",
            ticket.get("issue_type") or "",
            ticket.get("severity") or "",
            ticket.get("description") or "",
            ticket.get("resolution_notes") or "",
            ticket.get("status") or "",
            ticket.get("reported_by") or "",
            ticket.get("resolved_by") or "",
        ]).lower()
        score = sum(1 for w in q_words if w in text)
        if score > 0:
            scored.append((score, ticket))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [t for _, t in scored[:limit]]


def find_relevant_employees(question, limit=3):
    """Search for employees relevant to the question."""
    if supabase is None:
        return []
    try:
        res = supabase.table("users") \
            .select("employee_id,name,email,department,designation,role") \
            .execute()
        employees = res.data or []
    except Exception:
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)
    if not q_words:
        return []

    scored = []
    for emp in employees:
        text = " ".join([
            emp.get("employee_id") or "",
            emp.get("name") or "",
            emp.get("email") or "",
            emp.get("department") or "",
            emp.get("designation") or "",
            emp.get("role") or "",
        ]).lower()
        score = sum(1 for w in q_words if w in text)
        if score > 0:
            scored.append((score, emp))

    scored.sort(key=lambda x: x[0], reverse=True)
    return [e for _, e in scored[:limit]]


def find_relevant_approvals(question, limit=4):
    """
    Search document_approvals (Pending / Approved / Rejected) so PlantMind AI
    knows about files that are still in review or already decided - not just
    documents that made it all the way into the published Knowledge Base.
    This is what lets it answer things like "what's the status of the VPN
    policy upload" or "who submitted the onboarding SOP update".
    """
    if supabase is None:
        return []
    try:
        rows = supabase.table("document_approvals") \
            .select("id,doc_name,status,submitted_by,decided_by,decided_at,"
                    "old_version,proposed_version,new_content,ai_score,ai_reasoning,"
                    "file_name,created_at") \
            .order("created_at", desc=True).limit(300).execute().data or []
    except Exception as e:
        print("[WARN] find_relevant_approvals failed:", e)
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)
    if not q_words:
        return []

    scored = []
    for row in rows:
        text = " ".join([
            row.get("doc_name") or "", row.get("status") or "",
            row.get("submitted_by") or "", row.get("decided_by") or "",
            row.get("file_name") or "",
            (row.get("new_content") or "")[:500], row.get("ai_reasoning") or "",
        ]).lower()
        score = sum(1 for w in q_words if w in text)
        if score:
            scored.append((score, row))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [r for _, r in scored[:limit]]


def find_relevant_history(question, limit=4):
    """
    Search document_history so PlantMind AI can answer "who changed this
    document and what did the previous version say" - real, per-version
    provenance rather than just the latest published text.
    """
    if supabase is None:
        return []
    try:
        rows = supabase.table("document_history") \
            .select("id,document_id,version,updated_by,changes,content_snapshot,created_at") \
            .order("created_at", desc=True).limit(300).execute().data or []
    except Exception as e:
        print("[WARN] find_relevant_history failed:", e)
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)
    if not q_words:
        return []

    scored = []
    for row in rows:
        text = " ".join([
            row.get("version") or "", row.get("updated_by") or "",
            row.get("changes") or "", (row.get("content_snapshot") or "")[:300],
        ]).lower()
        score = sum(1 for w in q_words if w in text)
        if score:
            scored.append((score, row))
    scored.sort(key=lambda x: x[0], reverse=True)
    top = [r for _, r in scored[:limit]]

    # Attach the parent document's name so the context reads naturally
    # instead of just showing a raw document_id.
    if top:
        doc_ids = list({r["document_id"] for r in top if r.get("document_id")})
        if doc_ids:
            try:
                docs = supabase.table("documents").select("id,name").in_("id", doc_ids).execute().data or []
                name_map = {d["id"]: d["name"] for d in docs}
                for r in top:
                    r["doc_name"] = name_map.get(r.get("document_id"), "Unknown document")
            except Exception:
                for r in top:
                    r["doc_name"] = "Unknown document"
        else:
            for r in top:
                r["doc_name"] = "Unknown document"
    return top


def find_relevant_activity(question, limit=5):
    """
    Search activity_log for "who did what, when" - approvals, rejections,
    incident resolutions, uploads, deletions, batch automation runs - the
    full audit trail behind every dashboard number.
    """
    if supabase is None:
        return []
    try:
        rows = supabase.table("activity_log") \
            .select("actor,action,target,details,created_at") \
            .order("created_at", desc=True).limit(400).execute().data or []
    except Exception as e:
        print("[WARN] find_relevant_activity failed:", e)
        return []

    q_words = set(w.lower() for w in question.split() if len(w) > 2)
    if not q_words:
        return []

    scored = []
    for row in rows:
        text = " ".join([
            row.get("actor") or "", row.get("action") or "",
            row.get("target") or "", row.get("details") or "",
        ]).lower()
        score = sum(1 for w in q_words if w in text)
        if score:
            scored.append((score, row))
    scored.sort(key=lambda x: x[0], reverse=True)
    return [r for _, r in scored[:limit]]


@app.route("/api/ai/ask", methods=["POST"])
def ai_ask():
    if groq_client is None:
        return jsonify({"error": "Groq is not configured. Add GROQ_API_KEY to .env"}), 500

    data     = request.get_json(force=True) or {}
    question = (data.get("question") or "").strip()
    mode     = data.get("mode", "general")
    asked_by = data.get("asked_by", "Unknown")
    if not question:
        return jsonify({"error": "question is required"}), 400

    # Pull live, grounded context from every real data source in Supabase.
    relevant_docs = find_relevant_documents(question, limit=4)
    relevant_incidents = find_relevant_incidents(question, limit=3)
    relevant_employees = find_relevant_employees(question, limit=3)
    relevant_approvals = find_relevant_approvals(question, limit=4)
    relevant_history = find_relevant_history(question, limit=4)
    relevant_activity = find_relevant_activity(question, limit=5)

    for d in relevant_docs:
        d["content"] = ensure_doc_content(d)

    context_sections = []

    # Documents (Active AND Draft/Updated)
    for d in relevant_docs:
        context_sections.append(
            f"[Document: {d['name']} ({d.get('version','v1.0')}) — status: {d.get('status','Active')}]\n"
            f"{(d.get('content') or d.get('summary') or '')[:6000]}"
        )

    # Incidents (open + resolved)
    for ticket in relevant_incidents:
        context_sections.append(
            f"[Incident: {ticket.get('ticket_id', 'INC')}]\n"
            f"Status: {ticket.get('status', 'Open')}\n"
            f"System: {ticket.get('system_name', 'Unknown')}\n"
            f"Issue Type: {ticket.get('issue_type', 'Unknown')}\n"
            f"Severity: {ticket.get('severity', 'Unknown')}\n"
            f"Reported by: {ticket.get('reported_by', 'Unknown')} on {ticket.get('created_at','—')}\n"
            f"Description: {ticket.get('description', '')}\n"
            f"Resolved by: {ticket.get('resolved_by') or 'not yet resolved'}"
            f"{(' on ' + ticket.get('resolved_at')) if ticket.get('resolved_at') else ''}\n"
            f"Resolution Notes: {ticket.get('resolution_notes') or 'Pending'}"
        )

    # Approvals queue — Pending / Approved / Rejected requests, with AI score
    for a in relevant_approvals:
        context_sections.append(
            f"[Approval Request: {a.get('doc_name')} — status: {a.get('status')}]\n"
            f"Submitted by {a.get('submitted_by')} on {a.get('created_at')} "
            f"({a.get('old_version')} -> {a.get('proposed_version')})\n"
            f"Attached file: {a.get('file_name') or 'none / text only'}\n"
            f"AI match score: {a.get('ai_score')} — {a.get('ai_reasoning') or 'no reasoning recorded'}\n"
            f"Decided by: {a.get('decided_by') or 'not yet decided'} on {a.get('decided_at') or '—'}\n"
            f"Proposed content excerpt: {(a.get('new_content') or '')[:1500]}"
        )

    # Version history — real provenance per document
    for h in relevant_history:
        context_sections.append(
            f"[Version History: {h.get('doc_name')} — {h.get('version')}]\n"
            f"Changed by {h.get('updated_by')} on {h.get('created_at')}\n"
            f"Change notes: {h.get('changes')}\n"
            f"Content at that version (excerpt): {(h.get('content_snapshot') or '')[:1000]}"
        )

    # Activity log — the full audit trail
    for act in relevant_activity:
        context_sections.append(
            f"[Activity: {act.get('actor')} {act.get('action')} on {act.get('target')} "
            f"at {act.get('created_at')}]\n{act.get('details') or ''}"
        )

    # Employee directory
    for emp in relevant_employees:
        context_sections.append(
            f"[Employee: {emp.get('name', 'Unknown')} ({emp.get('employee_id', 'ID')})]\n"
            f"Department: {emp.get('department', 'Unknown')}\n"
            f"Designation: {emp.get('designation', emp.get('role', 'Unknown'))}\n"
            f"Email: {emp.get('email', 'Unknown')}\n"
            f"Role: {emp.get('role', 'Unknown')}"
        )

    context_text = "\n\n".join(context_sections) or "No matching internal records were found."

    instruction = MODE_INSTRUCTIONS.get(mode, MODE_INSTRUCTIONS["general"])

    system_prompt = f"""You are PlantMind AI, the internal knowledge assistant for OnlyTech.
{instruction}

You have access to real, live data pulled directly from OnlyTech's systems, including:
published and draft Knowledge Base documents, the document approval queue (pending,
approved, and rejected uploads with AI match scores), full version history for every
document, IT incident tickets (open and resolved), the system activity log (who did
what and when), and the employee directory.

Use ONLY the context below when it is relevant. If the context does not cover the
question, say so plainly and give general best-practice guidance instead. When asked
about status, history, or "who did X", answer from the specific records in the
context (submitter, decider, timestamps, version numbers) rather than guessing.

--- CONTEXT ---
{context_text}
--- END CONTEXT ---

Respond in clear, well-formatted plain text (short paragraphs / bullet points where helpful)."""

    try:
        chat_completion = groq_client.chat.completions.create(
            model=GROQ_MODEL_NAME,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user",   "content": question},
            ],
            temperature=0.4,
            max_tokens=1024,
        )
        answer_text = chat_completion.choices[0].message.content
    except Exception as e:
        return jsonify({"error": f"Groq request failed: {str(e)}"}), 502

    sources = (
        [{"type": "document", "id": d["id"], "name": d["name"], "version": d.get("version", "v1.0")}
         for d in relevant_docs]
        + [{"type": "incident", "id": t.get("id"),
            "name": f"{t.get('ticket_id', 'Incident')} · {t.get('system_name', 'Unknown')}",
            "version": t.get("status", "Open")}
           for t in relevant_incidents]
        + [{"type": "approval", "id": a.get("id"),
            "name": f"{a.get('doc_name')} (approval)", "version": a.get("status")}
           for a in relevant_approvals]
        + [{"type": "history", "id": h.get("id"),
            "name": f"{h.get('doc_name')} history", "version": h.get("version")}
           for h in relevant_history]
        + [{"type": "activity", "id": None,
            "name": f"{act.get('actor')}: {act.get('action')}", "version": ""}
           for act in relevant_activity]
        + [{"type": "employee", "id": e.get("employee_id"),
            "name": e.get("name", "Unknown"), "version": e.get("designation", e.get("role", ""))}
           for e in relevant_employees]
    )

    if supabase is not None:
        try:
            supabase.table("ai_query_logs").insert({
                "id":        str(uuid.uuid4()),
                "question":  question,
                "mode":      mode,
                "answer":    answer_text,
                "sources":   ", ".join(s["name"] for s in sources),
                "asked_by":  asked_by,
                "created_at": now_iso(),
            }).execute()
        except Exception:
            pass

    return jsonify({
        "answer":  answer_text,
        "mode":    mode,
        "sources": sources,
        "model":   GROQ_MODEL_NAME,
    }), 200

from flask import Flask, request, jsonify, Response, send_from_directory

@app.route("/")
def home():
    return send_from_directory(".", "admin.html")

@app.route("/employee")
def employee():
    return send_from_directory(".", "employee.html")
# --------------------------------------------------------------------------
if __name__ == "__main__":
    port = int(os.getenv("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)